import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os
import time


def fetch_chapter(chapter_url, doc):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36',
        'Referer': 'https://metanit.com/sharp/tutorial/'  # Изменен заголовок Referer
    }
    try:
        response = requests.get(chapter_url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # Извлечение заголовка страницы
        title = soup.find('title').text
        print(f"Обработка главы: {title}")  # Отладочный вывод
        doc.add_heading(title, level=1)

        # Проверяем, есть ли нужные элементы на странице
        has_content = False

        # Извлечение текстов из параграфов, заголовков и изображений
        for element in soup.find_all(['h3', 'p', 'pre', 'img']):
            if element.name == 'h3':
                doc.add_heading(element.text.strip(), level=2)  # Заголовки h3
                has_content = True
            elif element.name == 'p':
                doc.add_paragraph(element.text.strip())  # Параграфы
                has_content = True
            elif element.name == 'pre':
                # Обрабатываем код
                code_text = element.text.strip()
                doc.add_paragraph(code_text)  # Добавляем код как обычный параграф
                has_content = True
            elif element.name == 'img':
                img_url = element['src']
                print(f"Найдено изображение: {img_url}")  # Выводим URL изображения для отладки
                # Создаем полный URL для изображений
                if img_url.startswith('./'):
                    img_url = chapter_url.rsplit('/', 1)[0] + img_url[1:]  # Изменяем путь к изображению

                try:
                    img_response = requests.get(img_url, headers=headers)  # Добавляем заголовок Referer
                    img_response.raise_for_status()  # Проверяем статус ответа
                    # Сохраняем изображение во временный файл
                    img_file = 'temp_image.png'  # Имя временного файла для сохранения изображения
                    with open(img_file, 'wb') as f:
                        f.write(img_response.content)

                    # Добавляем изображение в документ на своем месте
                    doc.add_picture(img_file, width=Inches(5))  # Задаем ширину изображения

                    # Удаляем временный файл
                    os.remove(img_file)
                except Exception as img_error:
                    print(f"Не удалось загрузить изображение {img_url}: {img_error}")

        # Проверка, содержатся ли элементы на странице
        if not has_content:
            print(f"На странице {chapter_url} не найдено содержимого.")

    except Exception as e:
        print(f"Произошла ошибка при обработке {chapter_url}: {e}")


def fetch_all_chapters(start_chapter):
    base_url = "https://metanit.com/sharp/tutorial/"
    all_chapters = []

    # Составляем номера глав, начиная с введенного
    for i in range(1, 11):  # Предполагаем максимум 10 подглав
        chapter_number = f"{start_chapter}.{i}"
        chapter_url = f"{base_url}{chapter_number}.php"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36',
            'Referer': base_url
        }
        response = requests.get(chapter_url, headers=headers)

        print(f"Проверяем URL: {chapter_url}, Статус: {response.status_code}")  # Отладочный вывод

        if response.status_code == 200:
            all_chapters.append(chapter_number)
        else:
            print(f"Не удалось найти главу: {chapter_url}, Статус: {response.status_code}")  # Отладочный вывод
            break  # Выходим из цикла, если глава не найдена

    print(f"Найденные главы: {all_chapters}")  # Отладочный вывод
    return all_chapters


def main():
    # Ввод номера главы
    start_chapter = input("Введите номер главы (например, '1'): ")

    # Создаем документ
    doc = Document()

    # Получаем все главы, начиная с введенной
    chapters = fetch_all_chapters(start_chapter)

    # Обработка глав
    for chapter in chapters:
        chapter_url = f"https://metanit.com/sharp/tutorial/{chapter}.php"
        fetch_chapter(chapter_url, doc)
        time.sleep(1)  # Задержка между запросами

    # Сохраняем документ
    doc.save('all_chapters.docx')
    print("Документ сохранён как all_chapters.docx")


if __name__ == "__main__":
    main()
