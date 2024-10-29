import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os

def fetch_chapter(chapter_url, chapter):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36',
        'Referer': 'https://metanit.com/'  # Добавляем заголовок Referer
    }
    try:
        response = requests.get(chapter_url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')
        doc = Document()

        # Извлечение заголовка страницы
        title = soup.find('title').text
        doc.add_heading(title, level=1)

        # Извлечение текстов из параграфов, заголовков и изображений
        for element in soup.find_all(['h3', 'p', 'pre', 'img']):
            if element.name == 'h3':
                doc.add_heading(element.text.strip(), level=2)  # Заголовки h3
            elif element.name == 'p':
                doc.add_paragraph(element.text.strip())  # Параграфы
            elif element.name == 'pre':
                # Обрабатываем код
                code_text = element.text.strip()
                doc.add_paragraph(code_text)  # Добавляем код как обычный параграф
            elif element.name == 'img':
                img_url = element['src']
                print(f"Найдено изображение: {img_url}")  # Выводим URL изображения для отладки
                # Создаем полный URL для изображений
                if img_url.startswith('./'):
                    img_url = chapter_url.rsplit('/', 1)[0] + img_url[1:]  # Изменяем путь к изображению

                try:
                    img_response = requests.get(img_url, headers=headers)  # Добавляем заголовок Referer
                    img_response.raise_for_status()  # Проверяем статус ответа
                    # Сохраняем изображение во временный файл
                    img_file = 'temp_image.png'  # Имя временного файла для сохранения изображения
                    with open(img_file, 'wb') as f:
                        f.write(img_response.content)

                    # Добавляем изображение в документ на своем месте
                    doc.add_picture(img_file, width=Inches(5))  # Задаем ширину изображения

                    # Удаляем временный файл
                    os.remove(img_file)
                except Exception as img_error:
                    print(f"Не удалось загрузить изображение {img_url}: {img_error}")

        # Сохраняем документ с номером главы в названии
        file_name = f'chapter_{chapter}.docx'
        doc.save(file_name)
        print(f"Документ сохранён как {file_name}")
    except Exception as e:
        print(f"Произошла ошибка: {e}")

def main():
    base_url = "https://metanit.com/sharp/tutorial/"
    chapter = input("Введите номер главы (например, '1.1'): ")
    chapter_url = f"{base_url}{chapter}.php"

    fetch_chapter(chapter_url, chapter)

if __name__ == "__main__":
    main()
